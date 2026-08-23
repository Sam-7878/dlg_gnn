#!/usr/bin/env python3
"""
Extract THEIA specific attack details from TA51_Final_report_E5.docx
"""
import json
import re
from pathlib import Path
from docx import Document

DOCX_PATH = Path("/mnt/d/_Work/_data/DLG/DARPA-TC-THEIA/Ground_Truth/TA51_Final_report_E5.docx")

def main():
    if not DOCX_PATH.exists():
        print(f"Error: {DOCX_PATH} does not exist.")
        return

    doc = Document(str(DOCX_PATH))
    print(f"Loaded {DOCX_PATH.name} with {len(doc.paragraphs)} paragraphs.")

    # Search for headings and sections mentioning theia / ta51 / ta52 / attacks
    theia_sections = []
    current_section = None

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue

        style = p.style.name if p.style else ""
        if style.startswith("Heading"):
            if "theia" in text.lower() or "ta5" in text.lower() or "attack" in text.lower() or "day" in text.lower():
                current_section = {
                    "heading": text,
                    "style": style,
                    "para_idx": i,
                    "lines": []
                }
                theia_sections.append(current_section)
            elif current_section and style in ("Heading 1", "Heading 2"):
                # Started a new top-level heading not matching
                pass

        if current_section and len(current_section["lines"]) < 50:
            current_section["lines"].append(text)

    print(f"Found {len(theia_sections)} candidate sections.")
    for s in theia_sections[:20]:
        print(f"[{s['para_idx']}] {s['style']}: {s['heading']}")
        for line in s["lines"][:3]:
            print(f"    {line[:100]}")
        print()

if __name__ == "__main__":
    main()
